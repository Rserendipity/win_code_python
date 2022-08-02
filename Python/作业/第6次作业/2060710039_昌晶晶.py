from docx import Document

document = Document('./test2.docx')
# 读取段落
ps = [paragraph.text for paragraph in document.paragraphs]
# 读取表格
ts = [table for table in document.tables]
for t in ts:
    for row in t.rows:
        for cell in row.cells:
            ps.append(cell.text)
# 把字数和段落号转成字典
i = 1
d = dict()
for o in ps:
    if len(o) != 0:
        d[i] = len(o)
        i += 1

# 导入xlsx包
from openpyxl import Workbook
wb = Workbook()
temp = wb.active
temp.cell(row=1, column=1).value = '段落'
temp.cell(row=1, column=2).value = '字数'
i = 2
for key, val in d.items():
    temp.cell(row=i, column=1).value = key
    temp.cell(row=i, column=2).value = val
    i += 1
    # 用i来控制写入的行标
wb.save('./test1.xlsx')
