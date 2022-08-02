# -*- coding: utf-8 -*-
"""
Created on Sat May 21 11:28:22 2022

@author: Administrator
"""

from docx import Document

# https://zhuanlan.zhihu.com/p/146363527
# 打开文档
document = Document('d:/test2.docx')
# 读取标题、段落、列表内容
ps = [paragraph.text for paragraph in document.paragraphs]
for p in ps:
    print(p)
# 读取表格内容
ts = [table for table in document.tables]
for t in ts:
    for row in t.rows:
        for cell in row.cells:
            print(cell.text, end=' ')
        print()

# # https://blog.csdn.net/QQ6550523/article/details/106113528

from openpyxl import load_workbook, Workbook
# 打开xlsx文件
#wb = load_workbook('d:/test.xlsx')
wb = Workbook()
sheet = wb.active
sheet['A1'].value = 9999
sheet.cell(row=2, column=5).value = 'wonderful'
sheet.append((10, 20, 30, 40))
wb.save('d:/test2.xlsx')



## 遍历所有工作表的单元格
#for ws in wb.worksheets:
#    for row in ws.rows:
#        for cell in row:
#            try:
#                if dstStr in cell.value:
#                    return True
#            except:
#                pass
